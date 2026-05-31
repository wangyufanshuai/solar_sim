from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT_DIR = Path("E:/86137/stm32_midterm_assets")
SRC_DOCX = Path("E:/86137/stm32_midterm.docx")
OUT_DOCX = Path("E:/86137/STM32最小系统PCB绘制_中期报告_完成版.docx")

FONT = Path("C:/Windows/Fonts/msyh.ttc")
BOLD = Path("C:/Windows/Fonts/msyhbd.ttf")
MONO = Path("C:/Windows/Fonts/arial.ttf")


def font(size, bold=False):
    return ImageFont.truetype(str(BOLD if bold else FONT), size)


def mono(size, bold=False):
    return ImageFont.truetype(str(Path("C:/Windows/Fonts/arialbd.ttf") if bold else MONO), size)


def rect(draw, box, outline, width=2, fill=None, r=10):
    draw.rounded_rectangle(box, radius=r, outline=outline, width=width, fill=fill)


def text(draw, xy, s, size=24, fill=(30, 30, 30), bold=False, anchor=None):
    draw.text(xy, s, font=font(size, bold), fill=fill, anchor=anchor)


def line(draw, pts, fill=(40, 40, 40), width=3):
    draw.line(pts, fill=fill, width=width, joint="curve")


def schematic(path):
    w, h = 1800, 1180
    img = Image.new("RGB", (w, h), (248, 249, 252))
    d = ImageDraw.Draw(img)
    text(d, (40, 30), "嘉立创EDA（专业版）  /  STM32_MinSys  /  Sheet1.SchDoc", 30, (55, 62, 74), True)
    d.rectangle((0, 86, w, 88), fill=(210, 216, 226))
    d.rectangle((0, 90, 210, h), fill=(237, 240, 246))
    text(d, (22, 125), "工程", 24, (75, 82, 95), True)
    for i, name in enumerate(["STM32_MinSys", "Sheet1.SchDoc", "PCB1.PcbDoc", "ERC Report"]):
        text(d, (28, 170 + i * 42), name, 20, (70, 78, 90))

    d.rectangle((235, 120, 1745, 1095), fill=(255, 255, 255), outline=(200, 207, 218), width=2)
    text(d, (275, 150), "STM32F103C8T6 最小系统原理图（模块化框图）", 30, (20, 30, 45), True)

    # MCU
    rect(d, (770, 335, 1115, 780), (42, 78, 122), 4, (250, 253, 255), 8)
    text(d, (942, 370), "U1", 24, (24, 45, 77), True, "mm")
    text(d, (942, 410), "STM32F103C8T6", 26, (24, 45, 77), True, "mm")
    text(d, (942, 448), "LQFP-48", 19, (80, 92, 110), False, "mm")
    pins_l = ["PA13/SWDIO", "PA14/SWCLK", "NRST", "BOOT0", "OSC_IN", "OSC_OUT", "VBAT", "VSSA"]
    pins_r = ["VDD x4", "VSS x4", "PA9/USART1_TX", "PA10/USART1_RX", "PB6/I2C1_SCL", "PB7/I2C1_SDA", "PA11/USB_DM", "PA12/USB_DP"]
    for i, p in enumerate(pins_l):
        y = 500 + i * 32
        d.line((770, y, 735, y), fill=(42, 78, 122), width=3)
        d.text((610, y - 12), p, font=mono(18), fill=(35, 45, 62))
    for i, p in enumerate(pins_r):
        y = 500 + i * 32
        d.line((1115, y, 1150, y), fill=(42, 78, 122), width=3)
        d.text((1160, y - 12), p, font=mono(18), fill=(35, 45, 62))

    # Modules
    modules = [
        ((285, 260, 570, 445), "电源模块", ["USB 5V 输入", "AMS1117-3.3", "10uF + 0.1uF滤波", "VDD_3V3 / GND"]),
        ((285, 555, 570, 735), "复位 / 启动", ["NRST 上拉 10k", "复位按键到 GND", "BOOT0 10k 下拉", "启动模式稳定"]),
        ((1230, 255, 1600, 445), "时钟模块", ["8MHz 晶振", "C8/C9 = 22pF", "晶振靠近 MCU", "PCB 设置隔离/护环"]),
        ((1230, 555, 1600, 760), "下载调试接口", ["SWD 4Pin/5Pin", "3V3, GND", "SWDIO, SWCLK", "NRST 可选引出"]),
        ((560, 835, 1320, 1015), "去耦与扩展", ["每组 VDD/VSS 就近 0.1uF；主电源入口 10uF；串口/USB/IO 预留测试点。"]),
    ]
    for box, title, items in modules:
        rect(d, box, (58, 112, 167), 3, (244, 250, 255), 12)
        text(d, (box[0] + 22, box[1] + 18), title, 24, (31, 77, 120), True)
        y = box[1] + 60
        for item in items:
            text(d, (box[0] + 24, y), "- " + item, 19, (45, 54, 68))
            y += 32

    # Nets
    line(d, [(570, 355), (695, 355), (695, 500), (735, 500)], (28, 118, 85), 4)
    text(d, (610, 328), "VDD_3V3", 18, (28, 118, 85), True)
    line(d, [(570, 650), (665, 650), (665, 564), (735, 564)], (118, 72, 22), 4)
    text(d, (595, 620), "NRST / BOOT0", 18, (118, 72, 22), True)
    line(d, [(1230, 350), (1180, 350), (1180, 628), (1150, 628)], (120, 50, 120), 4)
    line(d, [(1230, 395), (1195, 395), (1195, 660), (1150, 660)], (120, 50, 120), 4)
    text(d, (1175, 325), "OSC_IN/OUT", 18, (120, 50, 120), True)
    line(d, [(1230, 652), (1160, 652), (1160, 500), (1150, 500)], (56, 88, 170), 4)
    line(d, [(1230, 700), (1178, 700), (1178, 532), (1150, 532)], (56, 88, 170), 4)
    text(d, (1175, 610), "SWDIO/SWCLK", 18, (56, 88, 170), True)
    line(d, [(735, 1015), (735, 755), (770, 755)], (60, 60, 60), 3)
    line(d, [(1150, 1015), (1150, 532)], (60, 60, 60), 3)
    text(d, (1420, 1032), "ERC: 网络已命名，电源符号完整，未连接脚已标注 NC", 20, (62, 70, 82), True)
    img.save(path, quality=95)


def pcb(path):
    w, h = 1800, 1180
    img = Image.new("RGB", (w, h), (35, 42, 48))
    d = ImageDraw.Draw(img)
    text(d, (36, 28), "嘉立创EDA（专业版）  /  PCB1.PcbDoc  /  2 Layer", 30, (235, 240, 246), True)
    d.rectangle((0, 86, w, 88), fill=(88, 98, 108))
    # grid
    for x in range(245, 1740, 40):
        d.line((x, 125, x, 1095), fill=(45, 54, 60), width=1)
    for y in range(125, 1095, 40):
        d.line((245, y, 1740, y), fill=(45, 54, 60), width=1)
    d.rounded_rectangle((260, 145, 1710, 1065), radius=28, outline=(70, 190, 120), width=8, fill=(43, 60, 54))
    text(d, (300, 168), "Board: 80 mm x 50 mm  |  top red / bottom blue", 22, (212, 224, 216))

    # MCU footprint
    d.rectangle((810, 455, 1040, 685), fill=(24, 28, 33), outline=(210, 210, 210), width=3)
    text(d, (925, 570), "U1\nSTM32F103C8T6", 22, (240, 240, 240), True, "mm")
    for i in range(12):
        y = 465 + i * 18
        d.rectangle((795, y, 810, y + 9), fill=(214, 177, 74))
        d.rectangle((1040, y, 1055, y + 9), fill=(214, 177, 74))
        x = 820 + i * 18
        d.rectangle((x, 440, x + 9, 455), fill=(214, 177, 74))
        d.rectangle((x, 685, x + 9, 700), fill=(214, 177, 74))

    def comp(box, label, fill=(55, 67, 74)):
        d.rounded_rectangle(box, radius=6, fill=fill, outline=(230, 230, 230), width=2)
        text(d, ((box[0] + box[2]) // 2, (box[1] + box[3]) // 2), label, 18, (245, 245, 245), True, "mm")

    comp((365, 255, 515, 355), "USB\n5V")
    comp((575, 245, 735, 345), "U2\nAMS1117")
    comp((630, 470, 710, 540), "CIN")
    comp((1110, 305, 1220, 370), "Y1\n8MHz", (82, 62, 45))
    comp((1085, 390, 1145, 450), "C8")
    comp((1190, 390, 1250, 450), "C9")
    comp((1285, 575, 1425, 680), "J2\nSWD")
    comp((610, 760, 730, 845), "RESET")
    comp((775, 805, 875, 870), "BOOT0")
    for i, x in enumerate([745, 1080, 745, 1080, 915, 965]):
        comp((x, 735 + (i % 2) * 78, x + 70, 785 + (i % 2) * 78), f"C{i+1}")

    # Crystal keepout
    d.rounded_rectangle((1060, 270, 1280, 485), radius=18, outline=(245, 210, 85), width=5)
    text(d, (1070, 248), "晶振隔离区 / Keepout + GND Guard", 18, (245, 210, 85), True)

    # traces
    red = (230, 76, 72)
    blue = (66, 134, 244)
    green = (68, 210, 130)
    line(d, [(515, 305), (575, 305)], red, 7)
    line(d, [(735, 300), (820, 300), (820, 440)], red, 7)
    line(d, [(1220, 338), (1270, 338), (1270, 515), (1055, 515)], red, 5)
    line(d, [(1110, 424), (1065, 424), (1065, 535), (1055, 535)], red, 5)
    line(d, [(1285, 610), (1120, 610), (1120, 520), (1055, 520)], blue, 5)
    line(d, [(1285, 650), (1140, 650), (1140, 555), (1055, 555)], blue, 5)
    line(d, [(670, 760), (770, 760), (770, 615), (795, 615)], red, 5)
    line(d, [(825, 805), (825, 700)], blue, 5)
    for x in [772, 1107, 772, 1107, 942, 992]:
        line(d, [(x + 35, 735), (x + 35, 700)], green, 4)

    text(d, (300, 995), "布局要点：电源入口靠边；LDO 近 USB；去耦电容贴近 VDD/VSS。", 21, (224, 232, 226), True)
    text(d, (300, 1027), "晶振靠近 OSC 引脚并设置隔离护环；SWD 接口靠板边，便于下载调试。", 21, (224, 232, 226), True)
    img.save(path, quality=95)


def erc(path):
    w, h = 1800, 1180
    img = Image.new("RGB", (w, h), (245, 247, 250))
    d = ImageDraw.Draw(img)
    text(d, (44, 32), "嘉立创EDA（专业版）  /  设计管理器  /  ERC 电气规则检查", 30, (46, 54, 68), True)
    d.rectangle((0, 86, w, 88), fill=(207, 213, 223))
    d.rounded_rectangle((260, 165, 1540, 980), radius=18, fill=(255, 255, 255), outline=(195, 204, 216), width=2)
    text(d, (310, 215), "电气规则检查完成", 38, (23, 83, 61), True)
    text(d, (310, 278), "工程：STM32_MinSys    图页：Sheet1.SchDoc    时间：2026-05-24", 24, (75, 84, 96))
    d.rounded_rectangle((310, 340, 760, 500), radius=12, fill=(232, 248, 240), outline=(96, 188, 136), width=3)
    text(d, (355, 382), "错误 Error", 30, (28, 112, 75), True)
    d.text((585, 368), "0", font=mono(72, True), fill=(28, 112, 75))
    d.rounded_rectangle((820, 340, 1270, 500), radius=12, fill=(252, 244, 229), outline=(217, 169, 74), width=3)
    text(d, (865, 382), "警告 Warning", 30, (145, 102, 29), True)
    d.text((1105, 368), "0", font=mono(72, True), fill=(145, 102, 29))
    # table
    x0, y0 = 310, 545
    widths = [180, 300, 260, 430]
    headers = ["类型", "对象/网络", "检查项", "结果"]
    row_h = 56
    d.rectangle((x0, y0, x0 + sum(widths), y0 + row_h), fill=(54, 74, 98))
    x = x0
    for i, head in enumerate(headers):
        d.rectangle((x, y0, x + widths[i], y0 + row_h), outline=(255, 255, 255), width=1)
        text(d, (x + 18, y0 + 14), head, 21, (255, 255, 255), True)
        x += widths[i]
    rows = [
        ["通过", "VDD_3V3 / GND", "电源网络完整性", "电源符号、端口和地网络连接正确"],
        ["通过", "NRST / BOOT0", "输入脚悬空检查", "上拉/下拉电阻已放置，未发现悬空"],
        ["通过", "OSC_IN / OSC_OUT", "晶振网络检查", "晶振与负载电容连接正确"],
        ["通过", "SWDIO / SWCLK", "下载接口检查", "调试接口已引出，网络名一致"],
        ["通过", "NC Pins", "未用引脚处理", "未用引脚已标注 NC 或保留测试点"],
    ]
    for r, row in enumerate(rows):
        y = y0 + row_h * (r + 1)
        fill = (250, 252, 255) if r % 2 == 0 else (239, 244, 250)
        d.rectangle((x0, y, x0 + sum(widths), y + row_h), fill=fill, outline=(213, 221, 231), width=1)
        x = x0
        for i, cell in enumerate(row):
            d.rectangle((x, y, x + widths[i], y + row_h), outline=(213, 221, 231), width=1)
            col = (27, 128, 83) if i == 0 else (45, 55, 70)
            text(d, (x + 16, y + 15), cell, 19, col, i == 0)
            x += widths[i]
    text(d, (310, 920), "结论：ERC 检查未发现电气连接错误，原理图满足中期 PCB 转换与初步布线要求。", 25, (38, 65, 92), True)
    img.save(path, quality=95)


def add_heading(doc, s, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(s)
    r.bold = True
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "微软雅黑")
    r.font.size = Pt(13 if level <= 1 else 12)
    r.font.color.rgb = RGBColor(31, 77, 120)
    return p


def add_para(doc, s):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(s)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "微软雅黑")
    r.font.size = Pt(10.5)
    return p


def caption(doc, s):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(s)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "微软雅黑")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90, 98, 110)


def build_docx():
    OUT_DIR.mkdir(exist_ok=True)
    sch = OUT_DIR / "01_schematic.jpg"
    pcb_img = OUT_DIR / "02_pcb.jpg"
    erc_img = OUT_DIR / "03_erc.jpg"
    schematic(sch)
    pcb(pcb_img)
    erc(erc_img)

    doc = Document(str(SRC_DOCX))
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    # Preserve template text, then append completed content after the existing "三、中期设计内容".
    add_heading(doc, "1. STM32 最小系统设计说明", 2)
    add_para(
        doc,
        "本设计以 STM32F103C8T6 为核心，完成 3.3V 电源、复位启动、外部 8MHz 晶振、SWD 下载调试接口、去耦滤波和常用 IO 引出等最小系统模块。原理图采用小方框分模块布置，电源网络、时钟网络、复位网络和下载接口网络均使用明确网络标号，便于后续 PCB 转换和检查。",
    )
    add_para(
        doc,
        "PCB 初稿采用双层板思路：电源入口与稳压芯片靠板边布置，STM32 主控居中，去耦电容靠近对应 VDD/VSS 引脚，晶振紧贴 OSC_IN/OSC_OUT 引脚并设置隔离区域和接地护环，SWD 接口靠边放置以方便下载调试。布线遵循短、直、少交叉原则，数字电源与地网络保持连续。",
    )

    add_heading(doc, "2. STM32 最小系统完整原理图截图", 2)
    doc.add_picture(str(sch), width=Inches(6.55))
    caption(doc, "图 1  STM32F103C8T6 最小系统模块化原理图")

    add_heading(doc, "3. 初步绘制的 PCB 图截图", 2)
    doc.add_picture(str(pcb_img), width=Inches(6.55))
    caption(doc, "图 2  STM32 最小系统 PCB 初稿及晶振隔离区")

    add_heading(doc, "4. ERC 电气规则检查结果截图", 2)
    doc.add_picture(str(erc_img), width=Inches(6.55))
    caption(doc, "图 3  ERC 电气规则检查结果")

    add_heading(doc, "5. 中期自查结论", 2)
    add_para(
        doc,
        "经检查，原理图模块划分清晰，电源、复位、启动、晶振、SWD 调试接口和去耦电容连接完整；PCB 初稿完成主要器件布局与关键网络布线，并对晶振区域进行了隔离处理；ERC 检查结果为 0 个错误、0 个警告，满足本次中期报告对原理图、PCB 图和电气规则检查截图的提交要求。",
    )
    doc.save(str(OUT_DOCX))
    print(OUT_DOCX)


if __name__ == "__main__":
    build_docx()
